from docx import Document
from docx2pdf import convert
import pandas as pd
import email
import smtplib
import ssl
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import os

df = pd.read_csv('list.csv')
path = 'Your Path'

address = 'Your Email Address'
password = 'Your Password'


def send_mail(name, receiver_email, file):
    mail_body = '''Hello, ''' + name + '''!
Your Message
    '''

    message = MIMEMultipart()
    message['From'] = address
    message['To'] = receiver_email
    message['Subject'] = 'Subject'

    message.attach(MIMEText(mail_body, 'plain'))

    filename = file

    with open(filename, "rb") as attachment:

        part = MIMEBase("application", "octet-stream")
        part.set_payload(attachment.read())

    encoders.encode_base64(part)

    part.add_header(
        "Content-Disposition",
        f"attachment; filename= {os.path.basename(file)}",
    )

    message.attach(part)
    text = message.as_string()

    context = ssl.create_default_context()
    with smtplib.SMTP_SSL("smtp.gmail.com", 465, context=context) as server:
        server.login(address, password)
        server.sendmail(address, receiver_email, text)

    print('Mail sent successfully!')


for name, receiver_email in df.itertuples(index=False):
    #print(name, receiver_email)
    document = Document('Template.docx')
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if 'Here' in run.text:
                            run.text = run.text.replace("Here", name)

    document.save(path + '/certificates-word/' + name + '.docx')
    convert(path + '/certificates-word/' + name + '.docx',
            path + '/certificates-pdf/' + name + '.pdf')

    file = path + '/certificates-pdf/' + name + '.pdf'
    print(file)
    send_mail(name, receiver_email, file)
